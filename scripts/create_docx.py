import subprocess
import sys
import os

def install(package):
    print(f"Instalando {package}...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install("python-docx")
    import docx

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Modificar estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Título Principal
    title = doc.add_paragraph()
    title_run = title.add_run("GUÍA DE AUTOMATIZACIÓN EDITORIAL CON IA")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Instrucciones paso a paso para clonar este proceso en futuros proyectos web")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")

    # Introducción
    doc.add_heading("1. Descripción del Sistema", level=1)
    doc.add_paragraph(
        "Este documento detalla el procedimiento automatizado para la creación en lote (batch) de artículos de "
        "blog optimizados para SEO y maquetados con Tailwind CSS, sin intervención manual. Este sistema "
        "conecta directamente la base de datos PostgreSQL del proyecto con la API de DeepSeek utilizando "
        "un script de automatización en Node.js."
    )

    # Requisitos previos
    doc.add_heading("2. Requisitos Previos", level=1)
    doc.add_paragraph("Para replicar este proceso en un nuevo sitio web, se necesita:")
    doc.add_paragraph("• Una cuenta activa en DeepSeek (con saldo en la plataforma de desarrollo).")
    doc.add_paragraph("• Una clave de API de DeepSeek directamente (puedes utilizar la clave actual: sk-177cec30fdd74db1ac9b23e204deade4).")
    doc.add_paragraph("• Una base de datos PostgreSQL conectada y con la tabla 'articles' creada y pre-poblada con los datos SEO básicos (título, keyword, slug, categoría, fecha, etc.) pero con el cuerpo de los artículos vacío ('').")

    # Configuración paso a paso
    doc.add_heading("3. Configuración Paso a Paso", level=1)
    
    doc.add_heading("Paso 3.1: Configurar el archivo de variables de entorno", level=2)
    doc.add_paragraph(
        "En la raíz del nuevo proyecto, crea o modifica el archivo '.env.local' para incluir la clave de la API de DeepSeek:"
    )
    p_code1 = doc.add_paragraph()
    p_code1.paragraph_format.left_indent = Inches(0.5)
    r_code1 = p_code1.add_run("DEEPSEEK_API_KEY=sk-177cec30fdd74db1ac9b23e204deade4\nDATABASE_URL=tu_conexion_postgresql_aqui")
    r_code1.font.name = 'Courier New'
    r_code1.font.size = Pt(10)

    doc.add_heading("Paso 3.2: Configurar la plantilla de redacción (template_general.md)", level=2)
    doc.add_paragraph(
        "Como plantilla de redacción se debe utilizar el archivo 'template_general.md' que se encuentra en la carpeta del proyecto. "
        "Este archivo contiene las instrucciones editoriales detalladas para la IA. Por ejemplo, restricciones de longitud "
        "(2.500-3.500+ palabras), la prohibición estricta de emojis, la densidad de keywords y la directiva de usar maquetación "
        "HTML y Tailwind CSS (como tablas y tarjetas de llamada)."
    )

    doc.add_heading("Paso 3.3: Crear el script de generación (generateArticles.mjs)", level=2)
    doc.add_paragraph(
        "Crea un script en la carpeta 'scripts/generateArticles.mjs' que lea los artículos vacíos de la base de datos, "
        "construya el prompt enviándolo a la API de DeepSeek ('https://api.deepseek.com/chat/completions' utilizando "
        "el modelo 'deepseek-chat') y guarde el HTML resultante de vuelta en la base de datos."
    )

    doc.add_heading("Paso 3.4: Ejecución", level=2)
    doc.add_paragraph("Ejecuta el script con el siguiente comando en tu terminal:")
    p_code2 = doc.add_paragraph()
    p_code2.paragraph_format.left_indent = Inches(0.5)
    r_code2 = p_code2.add_run("node scripts/generateArticles.mjs")
    r_code2.font.name = 'Courier New'
    r_code2.font.size = Pt(10)
    r_code2.bold = True

    # Prompt para pasarle al Agente de IA en el futuro
    doc.add_heading("4. Prompt para el Agente de IA", level=1)
    doc.add_paragraph(
        "Cuando inicies un nuevo proyecto web y quieras que yo (u otra IA) ejecute todo el proceso de forma autónoma, "
        "cópiale el siguiente texto literal:"
    )

    p_prompt = doc.add_paragraph()
    p_prompt.paragraph_format.left_indent = Inches(0.5)
    prompt_text = (
        "=== INSTRUCCIONES DE AUTOMATIZACIÓN EDITORIAL ===\n"
        "Hola. Necesito que automatices la generación de contenido para mi nueva web. Por favor, haz lo siguiente:\n"
        "1. Crea un script en la carpeta 'scripts/generateArticles.mjs' que:\n"
        "   - Cargue las variables de entorno de '.env.local' (incluyendo DEEPSEEK_API_KEY (clave a usar: sk-177cec30fdd74db1ac9b23e204deade4) y DATABASE_URL).\n"
        "   - Lea el archivo de plantilla llamado 'template_general.md' en la raíz.\n"
        "   - Obtenga todos los artículos de la base de datos PostgreSQL que tengan el campo 'content' vacío ('').\n"
        "   - Llame de forma secuencial a la API de DeepSeek (model: 'deepseek-chat', endpoint: 'https://api.deepseek.com/chat/completions') pasándole la plantilla y los metadatos de cada artículo (título, keyword, categoría).\n"
        "   - Guarde el resultado HTML devuelto en la columna 'content' de la tabla 'articles'.\n"
        "2. Usa el archivo 'template_general.md' en la raíz con directivas SEO estrictas: de 2500 a 3500+ palabras, estructurado en HTML con Tailwind CSS (tablas de comparativa, callouts), tono experto (EEAT) y sin emojis bajo ninguna circunstancia.\n"
        "3. Ejecuta el script localmente para procesar los artículos y actualizar la base de datos de producción/desarrollo."
    )
    r_prompt = p_prompt.add_run(prompt_text)
    r_prompt.font.name = 'Courier New'
    r_prompt.font.size = Pt(9.5)

    doc.add_paragraph("\n")
    doc.add_paragraph("Guardado y listo para futuras réplicas editoriales.")

    # Guardar documento
    output_filename = "api.docx"
    try:
        doc.save(output_filename)
        print(f"[OK] Documento '{output_filename}' creado con éxito en la raíz del proyecto.")
    except PermissionError:
        alternative = "api_new.docx"
        doc.save(alternative)
        print(f"[ERROR] No se pudo guardar en '{output_filename}' porque está abierto en Word.")
        print(f"[OK] Se ha guardado una copia en '{alternative}' en su lugar. Cierra Word y renombra el archivo.")

if __name__ == "__main__":
    create_document()
